from docxtpl import DocxTemplate
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import random
from pprint import pprint
import string

import subprocess
import os


def create_docx_from_data(template_name: str, context: dict, user_path: str):
    # Загружаем шаблон
    doc = Document(f"pdf_generator/templates/{template_name}.docx")

    letters = string.ascii_lowercase
    name = ''.join(random.choice(letters) for _ in range(8))

    # Загружаем шаблон
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for cell_num, cell in enumerate(cells):
                for context_key in context:
                    if context_key.startswith("char_"):
                        context_key_str = "{{"+context_key+"}}"    
                        if context_key_str in cell.text:
                            letters = list(context[context_key])
                            for i, ch in enumerate(letters):
                                if cell_num + i < len(row.cells):
                                    target_cell = row.cells[cell_num + i]
                                    target_cell.text = ch

                                    target_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    target_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                                    run = target_cell.paragraphs[0].runs[0]
                                    run.font.name = "Arial"
                                    run.font.size = Pt(12)

    doc.save(f"{user_path}/{name}.docx")

    #заполняем шаблон после заполнения клеток
    # Загружаем шаблон
    doc = DocxTemplate(f"{user_path}/{name}.docx")
    
    # Данные для подстановки
    # Рендерим документ
    doc.render(context)

    # Сохраняем результат
    doc.save(f"{user_path}/{name}.docx")
    return f"{user_path}/{name}.docx"


def convert_docx_to_pdf_libreoffice(input_docx_path, user_path=None):
    """
    Converts a DOCX file to a PDF using the LibreOffice command line.

    Args:
        input_docx_path (str): The path to the input .docx file.
        output_dir (str, optional): The directory where the PDF will be saved.
                                    If None, the PDF will be saved in the
                                    same directory as the input file.
    """
    if not os.path.exists(input_docx_path):
        pprint(f"Error: The file '{input_docx_path}' does not exist.")
        return

    if user_path is None:
        user_path = os.path.dirname(input_docx_path)

    temp_user_dir = "/tmp/libreoffice_user"
    os.makedirs(temp_user_dir, exist_ok=True)

    # The command to run LibreOffice in headless mode
    command = [
        "libreoffice",
        "--convert-to",
        "pdf",
        "--outdir",
        f"--env:UserInstallation=file://{temp_user_dir}",
        user_path,
        input_docx_path,
    ]

    try:
        # Run the command and capture output
        subprocess.run(command, check=True, capture_output=True, text=True)
        
        base_name = os.path.splitext(os.path.basename(input_docx_path))[0]
        pdf_path = os.path.join(user_path, f"{base_name}.pdf")

        
        pprint(f"Successfully converted '{input_docx_path}' to PDF in '{user_path}'.")

        return pdf_path
        
    except FileNotFoundError:
        pprint(
            "Error: LibreOffice executable not found. Please ensure it's installed and in your PATH."
        )
    except subprocess.CalledProcessError as e:
        pprint(f"An error occurred during conversion:")
        pprint(f"Command output: {e.stdout}")
        pprint(f"Command error: {e.stderr}")

def create_user_doc(user_path, template_name, context):
    user_path_docx = create_docx_from_data(
        template_name=template_name,
        context=context,
        user_path=user_path,
    )
    pprint(f"{user_path_docx}------------")
    pdf_path = convert_docx_to_pdf_libreoffice(input_docx_path=user_path_docx, user_path=user_path)
    return pdf_path

