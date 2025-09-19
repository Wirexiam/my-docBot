from aiogram import Router, F
from aiogram.types import CallbackQuery, Message, InlineKeyboardMarkup, InlineKeyboardButton, FSInputFile
from aiogram.fsm.context import FSMContext

from states.stay_prolong import StayProlong
from data_manager import SecureDataManager

# используем уже работающие модули паспорта
from handlers.components.passport_photo import start_new as passport_start_new_photo
from handlers.components.passport_manual import handle_passport_manual_start

stay_prolong_router = Router()
data_manager = SecureDataManager()


# ───────────────────────────── ВХОД ─────────────────────────────

@stay_prolong_router.callback_query(F.data == "stay_prolong")
async def sp_start(cb: CallbackQuery, state: FSMContext):
    """
    Точка входа: предлагается выбрать ввод паспорта (OCR/Manual).
    """
    await state.update_data(
        flow="stay_prolong",
        ocr_flow="sp",
        from_action=StayProlong.after_passport,  # возврат сюда после паспорта
        next_states=[],
        subflow=None
    )
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📸 Ввести паспорт по фото (OCR)", callback_data="sp_passport_photo")],
        [InlineKeyboardButton(text="⌨️ Ввести паспорт вручную", callback_data="sp_passport_manual")],
        [InlineKeyboardButton(text="🏠 В главное меню", callback_data="main_menu")]
    ])
    await state.set_state(StayProlong.start)
    await cb.message.edit_text(
        "Продление пребывания по браку / ребёнку / патенту\n\n"
        "Сначала внесём паспорт. Выберите способ:",
        reply_markup=kb
    )


@stay_prolong_router.callback_query(F.data == "sp_passport_photo")
async def sp_passport_photo(cb: CallbackQuery, state: FSMContext):
    """
    Запуск OCR ветки паспорта с нашими маркерами (ocr_flow=sp, from_action=StayProlong.after_passport).
    """
    await passport_start_new_photo(cb, state)


@stay_prolong_router.callback_query(F.data == "sp_passport_manual")
async def sp_passport_manual(cb: CallbackQuery, state: FSMContext):
    """
    Запуск ручного ввода паспорта с нашими маркерами (ocr_flow=sp, from_action=StayProlong.after_passport).
    """
    fake_cb = cb.model_copy(update={"data": "passport_new_manual_start"})
    await handle_passport_manual_start(fake_cb, state)


# ────────────────── возврат из паспорта (кнопка sp_after_passport) ──────────────────

@stay_prolong_router.callback_query(F.data == "sp_after_passport")
async def sp_after_passport(cb: CallbackQuery, state: FSMContext):
    """
    Пользователь подтвердил паспорт — выбираем основание продления.
    """
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="💍 По браку", callback_data="sp_basis_marriage")],
        [InlineKeyboardButton(text="👶 По ребёнку", callback_data="sp_basis_child")],
        [InlineKeyboardButton(text="📑 По патенту", callback_data="sp_basis_patent")],
        [InlineKeyboardButton(text="🔙 Назад к способу ввода паспорта", callback_data="stay_prolong")]
    ])
    await state.set_state(StayProlong.select_basis)
    await cb.message.edit_text(
        "Выберите основание продления пребывания:",
        reply_markup=kb
    )


# ───────────────────────────── БРАК ─────────────────────────────

@stay_prolong_router.callback_query(F.data == "sp_basis_marriage")
async def sp_basis_marriage(cb: CallbackQuery, state: FSMContext):
    await state.update_data(subflow="marriage", marriage_data={})
    await state.set_state(StayProlong.marriage_spouse_full_name)
    await cb.message.edit_text("Введите ФИО супруга(и) (кириллицей):")


@stay_prolong_router.message(StayProlong.marriage_spouse_full_name)
async def sp_marriage_spouse_full_name(msg: Message, state: FSMContext):
    sd = await state.get_data()
    m = dict(sd.get("marriage_data") or {})
    m["spouse_full_name"] = (msg.text or "").strip()
    await state.update_data(marriage_data=m)
    await state.set_state(StayProlong.marriage_spouse_birth_date)
    await msg.answer("Введите дату рождения супруга(и) в формате ДД.ММ.ГГГГ:")


@stay_prolong_router.message(StayProlong.marriage_spouse_birth_date)
async def sp_marriage_spouse_birth_date(msg: Message, state: FSMContext):
    sd = await state.get_data()
    m = dict(sd.get("marriage_data") or {})
    m["spouse_birth_date"] = (msg.text or "").strip()
    await state.update_data(marriage_data=m)
    await state.set_state(StayProlong.marriage_spouse_citizenship)
    await msg.answer("Укажите гражданство супруга(и) (кириллицей):")


@stay_prolong_router.message(StayProlong.marriage_spouse_citizenship)
async def sp_marriage_spouse_citizenship(msg: Message, state: FSMContext):
    sd = await state.get_data()
    m = dict(sd.get("marriage_data") or {})
    m["spouse_citizenship"] = (msg.text or "").strip()
    await state.update_data(marriage_data=m)
    await state.set_state(StayProlong.marriage_cert_number)
    await msg.answer("Введите номер свидетельства о браке:")


@stay_prolong_router.message(StayProlong.marriage_cert_number)
async def sp_marriage_cert_number(msg: Message, state: FSMContext):
    sd = await state.get_data()
    m = dict(sd.get("marriage_data") or {})
    m["cert_number"] = (msg.text or "").strip()
    await state.update_data(marriage_data=m)
    await state.set_state(StayProlong.marriage_cert_date)
    await msg.answer("Введите дату выдачи свидетельства о браке (ДД.ММ.ГГГГ):")


@stay_prolong_router.message(StayProlong.marriage_cert_date)
async def sp_marriage_cert_date(msg: Message, state: FSMContext):
    sd = await state.get_data()
    m = dict(sd.get("marriage_data") or {})
    m["cert_date"] = (msg.text or "").strip()
    await state.update_data(marriage_data=m)
    await state.set_state(StayProlong.marriage_cert_issued_by)
    await msg.answer("Укажите КЕМ выдано свидетельство о браке:")


@stay_prolong_router.message(StayProlong.marriage_cert_issued_by)
async def sp_marriage_cert_issued_by(msg: Message, state: FSMContext):
    sd = await state.get_data()
    m = dict(sd.get("marriage_data") or {})
    m["cert_issued_by"] = (msg.text or "").strip()
    await state.update_data(marriage_data=m)
    await ask_address(msg, state)


# ───────────────────────────── РЕБЁНОК ─────────────────────────────

@stay_prolong_router.callback_query(F.data == "sp_basis_child")
async def sp_basis_child(cb: CallbackQuery, state: FSMContext):
    await state.update_data(subflow="child", child_data={})
    await state.set_state(StayProlong.child_full_name)
    await cb.message.edit_text("Введите ФИО ребёнка (кириллицей):")


@stay_prolong_router.message(StayProlong.child_full_name)
async def sp_child_full_name(msg: Message, state: FSMContext):
    sd = await state.get_data()
    d = dict(sd.get("child_data") or {})
    d["child_full_name"] = (msg.text or "").strip()
    await state.update_data(child_data=d)
    await state.set_state(StayProlong.child_birth_date)
    await msg.answer("Введите дату рождения ребёнка (ДД.ММ.ГГГГ):")


@stay_prolong_router.message(StayProlong.child_birth_date)
async def sp_child_birth_date(msg: Message, state: FSMContext):
    sd = await state.get_data()
    d = dict(sd.get("child_data") or {})
    d["child_birth_date"] = (msg.text or "").strip()
    await state.update_data(child_data=d)
    await state.set_state(StayProlong.child_citizenship)
    await msg.answer("Укажите гражданство ребёнка (кириллицей):")


@stay_prolong_router.message(StayProlong.child_citizenship)
async def sp_child_citizenship(msg: Message, state: FSMContext):
    sd = await state.get_data()
    d = dict(sd.get("child_data") or {})
    d["child_citizenship"] = (msg.text or "").strip()
    await state.update_data(child_data=d)
    await state.set_state(StayProlong.child_cert_number)
    await msg.answer("Введите номер свидетельства о рождении:")


@stay_prolong_router.message(StayProlong.child_cert_number)
async def sp_child_cert_number(msg: Message, state: FSMContext):
    sd = await state.get_data()
    d = dict(sd.get("child_data") or {})
    d["child_cert_number"] = (msg.text or "").strip()
    await state.update_data(child_data=d)
    await state.set_state(StayProlong.child_cert_date)
    await msg.answer("Введите дату выдачи свидетельства о рождении (ДД.ММ.ГГГГ):")


@stay_prolong_router.message(StayProlong.child_cert_date)
async def sp_child_cert_date(msg: Message, state: FSMContext):
    sd = await state.get_data()
    d = dict(sd.get("child_data") or {})
    d["child_cert_date"] = (msg.text or "").strip()
    await state.update_data(child_data=d)
    await state.set_state(StayProlong.child_cert_issued_by)
    await msg.answer("Укажите КЕМ выдано свидетельство о рождении:")


@stay_prolong_router.message(StayProlong.child_cert_issued_by)
async def sp_child_cert_issued_by(msg: Message, state: FSMContext):
    sd = await state.get_data()
    d = dict(sd.get("child_data") or {})
    d["child_cert_issued_by"] = (msg.text or "").strip()
    await state.update_data(child_data=d)
    await ask_address(msg, state)


# ───────────────────────────── ПАТЕНТ ─────────────────────────────

@stay_prolong_router.callback_query(F.data == "sp_basis_patent")
async def sp_basis_patent(cb: CallbackQuery, state: FSMContext):
    await state.update_data(subflow="patent", patent_data={})
    await state.set_state(StayProlong.patent_number)
    await cb.message.edit_text("Введите номер патента (пример: 78-123456789):")


@stay_prolong_router.message(StayProlong.patent_number)
async def sp_patent_number(msg: Message, state: FSMContext):
    sd = await state.get_data()
    p = dict(sd.get("patent_data") or {})
    p["patent_number"] = (msg.text or "").strip()
    await state.update_data(patent_data=p)
    await state.set_state(StayProlong.patent_issue_date)
    await msg.answer("Введите дату выдачи патента (ДД.ММ.ГГГГ):")


@stay_prolong_router.message(StayProlong.patent_issue_date)
async def sp_patent_issue_date(msg: Message, state: FSMContext):
    sd = await state.get_data()
    p = dict(sd.get("patent_data") or {})
    p["patent_issue_date"] = (msg.text or "").strip()
    await state.update_data(patent_data=p)
    await state.set_state(StayProlong.patent_issued_by)
    await msg.answer("Укажите, кем выдан патент:")


@stay_prolong_router.message(StayProlong.patent_issued_by)
async def sp_patent_issued_by(msg: Message, state: FSMContext):
    sd = await state.get_data()
    p = dict(sd.get("patent_data") or {})
    p["patent_issued_by"] = (msg.text or "").strip()
    await state.update_data(patent_data=p)
    await state.set_state(StayProlong.patent_profession)
    await msg.answer("Укажите профессию (как в патенте):")


@stay_prolong_router.message(StayProlong.patent_profession)
async def sp_patent_profession(msg: Message, state: FSMContext):
    sd = await state.get_data()
    p = dict(sd.get("patent_data") or {})
    p["profession"] = (msg.text or "").strip()
    await state.update_data(patent_data=p)
    await state.set_state(StayProlong.patent_employer_address)
    await msg.answer("Введите юридический адрес работодателя (индекс, город, улица, дом):")


@stay_prolong_router.message(StayProlong.patent_employer_address)
async def sp_patent_employer_address(msg: Message, state: FSMContext):
    sd = await state.get_data()
    p = dict(sd.get("patent_data") or {})
    p["employer_address"] = (msg.text or "").strip()
    await state.update_data(patent_data=p)
    await state.set_state(StayProlong.patent_inn)
    await msg.answer("Введите ИНН работодателя (10 или 12 цифр):")


@stay_prolong_router.message(StayProlong.patent_inn)
async def sp_patent_inn(msg: Message, state: FSMContext):
    sd = await state.get_data()
    p = dict(sd.get("patent_data") or {})
    p["employer_inn"] = (msg.text or "").strip()
    await state.update_data(patent_data=p)
    # ДМС — опционально, спросим коротко
    await state.set_state(StayProlong.patent_dms_number)
    await msg.answer("Номер полиса ДМС (или '-' если нет):")


@stay_prolong_router.message(StayProlong.patent_dms_number)
async def sp_patent_dms_number(msg: Message, state: FSMContext):
    sd = await state.get_data()
    p = dict(sd.get("patent_data") or {})
    p["dms_number"] = (msg.text or "").strip()
    await state.update_data(patent_data=p)
    await state.set_state(StayProlong.patent_dms_company)
    await msg.answer("Страховая компания (или '-' если нет):")


@stay_prolong_router.message(StayProlong.patent_dms_company)
async def sp_patent_dms_company(msg: Message, state: FSMContext):
    sd = await state.get_data()
    p = dict(sd.get("patent_data") or {})
    p["dms_company"] = (msg.text or "").strip()
    await state.update_data(patent_data=p)
    await state.set_state(StayProlong.patent_dms_period)
    await msg.answer("Срок ДМС (например: с 15.06.2025 по 14.06.2026) или '-' :")


@stay_prolong_router.message(StayProlong.patent_dms_period)
async def sp_patent_dms_period(msg: Message, state: FSMContext):
    sd = await state.get_data()
    p = dict(sd.get("patent_data") or {})
    p["dms_period"] = (msg.text or "").strip()
    await state.update_data(patent_data=p)
    await ask_address(msg, state)


# ───────────────────────────── АДРЕС/ТЕЛЕФОН ─────────────────────────────
# Для простоты делаем свои промпты. Если захочешь — легко заменить на ваш ask_live_adress.

async def ask_address(msg_or_cb, state: FSMContext):
    if isinstance(msg_or_cb, Message):
        message = msg_or_cb
    else:
        message = msg_or_cb.message
    await state.set_state(StayProlong.address)
    await message.answer(
        "Введите адрес фактического проживания (одной строкой):\n"
        "город, улица, дом, корпус/строение (если есть), квартира"
    )


@stay_prolong_router.message(StayProlong.address)
async def sp_address(msg: Message, state: FSMContext):
    await state.update_data(sp_address=(msg.text or "").strip())
    await state.set_state(StayProlong.phone)
    await msg.answer("Введите номер телефона в формате 79XXXXXXXXX:")


@stay_prolong_router.message(StayProlong.phone)
async def sp_phone(msg: Message, state: FSMContext):
    await state.update_data(sp_phone=(msg.text or "").strip())
    await show_confirm(msg, state)


# ───────────────────────────── СВОДКА + ГЕНЕРАЦИЯ ─────────────────────────────

async def show_confirm(msg_or_cb, state: FSMContext):
    if isinstance(msg_or_cb, Message):
        message = msg_or_cb
    else:
        message = msg_or_cb.message

    sd = await state.get_data()
    pd = sd.get("passport_data") or {}
    old_pd = sd.get("old_passport_data") or {}
    basis = sd.get("subflow")
    addr = sd.get("sp_address", "—")
    phone = sd.get("sp_phone", "—")

    def val(d, k):
        v = (d.get(k) or "").strip()
        return v if v else "—"

    text = (
        "📋 Подтверждение данных\n\n"
        f"👤 ФИО: {val(pd, 'full_name')}\n"
        f"📄 Паспорт: {val(pd, 'passport_serial_number')}, "
        f"{val(pd, 'passport_issue_date')}, {val(pd, 'passport_issue_place')}, {val(pd, 'passport_expiry_date')}\n"
    )
    if old_pd:
        text += f"📄 Старый паспорт: {val(old_pd, 'passport_serial_number')} / {val(old_pd, 'passport_issue_date')}\n"

    if basis == "marriage":
        m = sd.get("marriage_data") or {}
        text += (
            "\n💍 Основание: Брак\n"
            f"Супруг(а): {m.get('spouse_full_name', '—')} "
            f"({m.get('spouse_birth_date', '—')}, {m.get('spouse_citizenship', '—')})\n"
            f"Свидетельство о браке: №{m.get('cert_number', '—')} от {m.get('cert_date', '—')}, "
            f"{m.get('cert_issued_by', '—')}\n"
        )
        filename = "Заявление_о_продлениеи_по_браку.docx"
    elif basis == "child":
        d = sd.get("child_data") or {}
        text += (
            "\n👶 Основание: Ребёнок\n"
            f"Ребёнок: {d.get('child_full_name', '—')} "
            f"({d.get('child_birth_date', '—')}, {d.get('child_citizenship', '—')})\n"
            f"Свидетельство о рождении: №{d.get('child_cert_number', '—')} от {d.get('child_cert_date', '—')}, "
            f"{d.get('child_cert_issued_by', '—')}\n"
        )
        filename = "Заявление_о_продлении_по_ребенку.docx"
    else:
        p = sd.get("patent_data") or {}
        text += (
            "\n📑 Основание: Патент\n"
            f"Патент: {p.get('patent_number', '—')} от {p.get('patent_issue_date', '—')}, "
            f"{p.get('patent_issued_by', '—')}\n"
            f"Профессия: {p.get('profession', '—')}\n"
            f"Работодатель: {p.get('employer_address', '—')} | ИНН: {p.get('employer_inn', '—')}\n"
            f"ДМС: {p.get('dms_number', '—')}, {p.get('dms_company', '—')}, {p.get('dms_period', '—')}\n"
        )
        filename = "Заявление_о_продление_по_патенту.docx"

    text += f"\n🏠 Адрес: {addr}\n📞 Телефон: {phone}"

    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Всё верно — сформировать документ", callback_data=f"sp_generate::{filename}")],
        [InlineKeyboardButton(text="✏ Изменить данные", callback_data="stay_prolong")]
    ])
    await state.set_state(StayProlong.confirm)
    await message.answer(text, reply_markup=kb)


@stay_prolong_router.callback_query(F.data.startswith("sp_generate::"))
async def sp_generate(cb: CallbackQuery, state: FSMContext):
    """
    Генерация DOCX «на месте» через python-docx, отправка файла пользователю.
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        await cb.message.answer("Не найден python-docx. Установите пакет и повторите.")
        return

    sd = await state.get_data()
    filename = cb.data.split("::", 1)[1]

    pd = sd.get("passport_data") or {}
    old_pd = sd.get("old_passport_data") or {}
    flow = sd.get("subflow")
    addr = sd.get("sp_address", "")
    phone = sd.get("sp_phone", "")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    doc.add_heading("Заявление о продлении пребывания", level=1)

    def line(text: str):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(11)

    line(f"ФИО: {pd.get('full_name', '')}")
    line(f"Паспорт: {pd.get('passport_serial_number', '')}, "
         f"{pd.get('passport_issue_date', '')}, "
         f"{pd.get('passport_issue_place', '')}, "
         f"{pd.get('passport_expiry_date', '')}")
    if old_pd:
        line(f"Старый паспорт: {old_pd.get('passport_serial_number', '')} / {old_pd.get('passport_issue_date', '')}")

    if flow == "marriage":
        m = sd.get("marriage_data") or {}
        doc.add_heading("Основание: Брак", level=2)
        line(f"Супруг(а): {m.get('spouse_full_name', '')} "
             f"({m.get('spouse_birth_date', '')}, {m.get('spouse_citizenship', '')})")
        line(f"Свидетельство о браке: №{m.get('cert_number', '')} от {m.get('cert_date', '')}, "
             f"{m.get('cert_issued_by', '')}")
    elif flow == "child":
        d = sd.get("child_data") or {}
        doc.add_heading("Основание: Ребёнок", level=2)
        line(f"Ребёнок: {d.get('child_full_name', '')} "
             f"({d.get('child_birth_date', '')}, {d.get('child_citizenship', '')})")
        line(f"Свидетельство о рождении: №{d.get('child_cert_number', '')} от {d.get('child_cert_date', '')}, "
             f"{d.get('child_cert_issued_by', '')}")
    else:
        p = sd.get("patent_data") or {}
        doc.add_heading("Основание: Патент", level=2)
        line(f"Патент: {p.get('patent_number', '')} от {p.get('patent_issue_date', '')}, "
             f"{p.get('patent_issued_by', '')}")
        line(f"Профессия: {p.get('profession', '')}")
        line(f"Работодатель: {p.get('employer_address', '')} | ИНН: {p.get('employer_inn', '')}")
        line(f"ДМС: {p.get('dms_number', '')}, {p.get('dms_company', '')}, {p.get('dms_period', '')}")

    doc.add_heading("Контакты", level=2)
    line(f"Адрес: {addr}")
    line(f"Телефон: {phone}")

    path = f"/tmp/{filename}"
    doc.save(path)

    await cb.message.answer_document(FSInputFile(path))
    await cb.answer("Готово ✅")

    # очищать state не будем — пусть юзер сможет отредактировать и сгенерировать заново при желании
